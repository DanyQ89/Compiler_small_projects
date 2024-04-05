from docx import Document


class CinemaHall:
    def __init__(self, name, rows, seats_per_row):
        self.name = name
        self.rows = rows
        self.seats_per_row = seats_per_row
        self.seat_matrix = [['O' for _ in range(seats_per_row)] for _ in range(rows)]

    def display_seats(self):
        for row in self.seat_matrix:
            print(" ".join(row))

    def is_seat_available(self, row, seat):
        return self.seat_matrix[row - 1][seat - 1] == 'O'

    def book_seat(self, row, seat):
        if self.is_seat_available(row, seat):
            self.seat_matrix[row - 1][seat - 1] = 'X'
            return True
        else:
            return False

    def get_available_seats(self):
        available_seats = []
        for i in range(self.rows):
            for j in range(self.seats_per_row):
                if self.seat_matrix[i][j] == 'O':
                    available_seats.append((i + 1, j + 1))
        return available_seats


class Cinema:
    def __init__(self, name):
        self.name = name
        self.halls = []

    def add_hall(self, hall):
        self.halls.append(hall)

    def list_halls(self):
        for hall in self.halls:
            print(hall.name)

    def get_next_showtime(self, movie_name):
        for hall in self.halls:
            for show in hall.shows:
                if show.movie_name == movie_name and show.has_available_seats():
                    return f"{show.start_time} in {self.name}, {hall.name}"
        return "No upcoming shows for the selected movie."

    def book_ticket(self, movie_name, row, seat):
        for hall in self.halls:
            for show in hall.shows:
                if show.movie_name == movie_name and show.has_available_seats():
                    if show.book_seat(row, seat):
                        return f"Ticket booked for {movie_name} at {show.start_time} in {self.name}, {hall.name}"
                    else:
                        return "Seat is already booked."
        return "No upcoming shows for the selected movie."

    def show_floor_plan(self, movie_name):
        for hall in self.halls:
            for show in hall.shows:
                if show.movie_name == movie_name:
                    print(f"Floor plan for {movie_name} at {show.start_time} in {self.name}, {hall.name}:")
                    show.display_seats()


class Show:
    def __init__(self, movie_name, start_time, duration):
        self.movie_name = movie_name
        self.start_time = start_time
        self.duration = duration
        self.hall = None
        self.seat_matrix = []

    def set_hall(self, hall):
        self.hall = hall
        self.seat_matrix = hall.seat_matrix

    def has_available_seats(self):
        return any('O' in row for row in self.seat_matrix)

    def book_seat(self, row, seat):
        return self.hall.book_seat(row, seat)


def save_to_docx(cinemas):
    doc = Document()
    doc.add_heading('Список кинотеатров, залов и сеансов', 0)

    for cinema in cinemas:
        doc.add_heading(f'Кинотеатр: {cinema.name}', level=1)

        for hall in cinema.halls:
            doc.add_heading(f'Зал: {hall.name}', level=2)
            doc.add_paragraph(f'Количество рядов: {hall.rows}')
            doc.add_paragraph(f'Мест в ряду: {hall.seats_per_row}')

            for show in hall.shows:
                doc.add_heading(f'Сеанс: {show.movie_name}', level=3)
                doc.add_paragraph(f'Начало: {show.start_time}')
                doc.add_paragraph(f'Длительность: {show.duration} минут')

    doc.save('cinemas.docx')
    print('Данные сохранены в файл cinemas.docx')
